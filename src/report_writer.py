from io import BytesIO
from typing import Dict, List

from docx import Document


def _add_paragraph_with_bold(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    if "**" not in text:
        p.add_run(text)
        return
    parts = text.split("**")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        if i % 2 == 1:
            run.bold = True


def _add_bullet_list(doc: Document, items: List[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_investor_report_docx(
    company: str,
    sections: Dict[str, str],
    highlights: List[str],
    achievement: str,
    funding_plan: str,
    recommendation: str,
    headings: List[str],
    include_recommendation: bool = True,
) -> bytes:
    doc = Document()
    doc.add_heading(f"{company} 투자자용 요약 및 추천", level=1)

    for heading in headings:
        doc.add_heading(heading, level=2)
        if heading.lower() == "highlights":
            _add_bullet_list(doc, highlights)
        elif heading.lower().startswith("achievement"):
            _add_paragraph_with_bold(doc, achievement)
        elif heading.lower().startswith("funding plan"):
            _add_paragraph_with_bold(doc, funding_plan)
        else:
            _add_paragraph_with_bold(doc, sections.get(heading, ""))

    if include_recommendation:
        doc.add_heading("Recommendation", level=2)
        _add_paragraph_with_bold(doc, recommendation)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_feedback_report_docx(company: str, feedback: Dict, total_score_100: float) -> bytes:
    doc = Document()
    doc.add_heading(f"{company} IR 상세 피드백: 총점 {round(total_score_100)}점", level=1)

    overall = feedback.get("overall_summary", "")
    if overall:
        doc.add_heading("종합 요약", level=2)
        _add_paragraph_with_bold(doc, overall)

    sections = feedback.get("sections", {})
    for name, info in sections.items():
        doc.add_heading(f"{name} (점수: {info.get('score_0_10', '')})", level=2)
        _add_paragraph_with_bold(doc, f"✅ 강점: {info.get('strengths', '')}")
        _add_paragraph_with_bold(doc, f"❌ 보완사항: {info.get('weaknesses', '')}")
        _add_paragraph_with_bold(doc, f"💡 보완 제안: {info.get('improvements', '')}")
        _add_paragraph_with_bold(doc, f"리스크/기대요소: {info.get('risks_expectations', '')}")

    priorities = feedback.get("priorities", "")
    if priorities:
        doc.add_heading("보완 우선순위", level=2)
        _add_paragraph_with_bold(doc, priorities)

    investor_strategy = feedback.get("investor_type_strategy", "")
    if investor_strategy:
        doc.add_heading("투자자 유형별 전략", level=2)
        _add_paragraph_with_bold(doc, investor_strategy)

    stage_guidelines = feedback.get("stage_guidelines", "")
    if stage_guidelines:
        doc.add_heading("성장단계별 가이드라인", level=2)
        _add_paragraph_with_bold(doc, stage_guidelines)

    pitch_faq = feedback.get("pitch_faq_strategy", "")
    if pitch_faq:
        doc.add_heading("피칭/FAQ 전략", level=2)
        _add_paragraph_with_bold(doc, pitch_faq)

    visual = feedback.get("visual_suggestions", "")
    if visual:
        doc.add_heading("시각적 보완 제안", level=2)
        _add_paragraph_with_bold(doc, visual)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
